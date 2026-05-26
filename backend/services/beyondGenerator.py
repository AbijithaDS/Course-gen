import sys
import os
import json
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def to_roman(num):
    try:
        n = int(num)
        if n <= 0:
            return str(num)
        roman_map = [
            (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
        ]
        result = ''
        for value, symbol in roman_map:
            while n >= value:
                result += symbol
                n -= value
        return result
    except:
        return str(num)

def parse_beyond_markdown(md_text):
    units = []
    # Match headers like "### Unit I: Title" or "### Unit 1: Title"
    matches = list(re.finditer(r"(?:^|\n)\s*###+\s*(Unit\s+[IVX\d]+[:\-\s].*?)(?=\n|$)", md_text, re.IGNORECASE))
    
    for i, match in enumerate(matches):
        header = match.group(1).strip()
        start = match.end()
        end = matches[i+1].start() if i + 1 < len(matches) else len(md_text)
        unit_body = md_text[start:end].strip()
        
        # Extract lines
        lines = unit_body.split('\n')
        topics = []
        for line in lines:
            line_str = line.strip()
            # Match line starts with numbers, e.g. "1. Topic Name - Description"
            # Separator can be dash, en-dash, em-dash, colon, or tilde
            m = re.match(r"^\s*(\d+)[\.\)]\s*(.*?)\s*[\–\-\—\:]\s*(.*)", line_str)
            if m:
                num = m.group(1).strip()
                name = m.group(2).strip()
                desc = m.group(3).strip()
                
                # Strip markdown asterisks or underscores
                name = re.sub(r"\*+|_+", "", name).strip()
                desc = re.sub(r"\*+|_+", "", desc).strip()
                
                topics.append({
                    "num": num,
                    "name": name,
                    "desc": desc
                })
            elif line_str and re.match(r"^\s*(\d+)[\.\)]\s*(.*)", line_str):
                # Fallback if no clean separator is found
                m_fallback = re.match(r"^\s*(\d+)[\.\)]\s*(.*)", line_str)
                num = m_fallback.group(1).strip()
                full_text = m_fallback.group(2).strip()
                full_text = re.sub(r"\*+|_+", "", full_text).strip()
                
                # Split by first double space or first period if any
                parts = full_text.split("  ", 1)
                if len(parts) > 1:
                    name = parts[0].strip()
                    desc = parts[1].strip()
                else:
                    name = full_text
                    desc = ""
                topics.append({
                    "num": num,
                    "name": name,
                    "desc": desc
                })
                
        # Clean header title
        title_parts = header.split(':', 1)
        unit_title = title_parts[1].strip() if len(title_parts) > 1 else header
        unit_title = re.sub(r"\*+|_+", "", unit_title).strip()
        
        units.append({
            "title": unit_title,
            "topics": topics
        })
        
    return units

def replace_text_in_runs(paragraphs, placeholder, replacement):
    for p in paragraphs:
        if placeholder in p.text:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)

def main():
    if len(sys.argv) < 4:
        print("Usage: python beyondGenerator.py <templatePath> <inputPath> <outputPath>")
        sys.exit(1)
        
    template_path = sys.argv[1]
    input_path = sys.argv[2]
    output_path = sys.argv[3]
    
    # 1. Load dynamic input payload
    with open(input_path, 'r', encoding='utf-8') as f:
        payload = json.load(f)
        
    subjectCode = payload.get("subjectCode", "")
    subjectName = payload.get("subjectName", "Academic Subject")
    staffName = payload.get("staffName", "Faculty member")
    departmentName = payload.get("departmentName", "Academic Department")
    year = payload.get("year", "")
    semester = payload.get("semester", "")
    regulation = payload.get("regulation", "")
    content = payload.get("content", "")
    
    # Format metadata
    yearRoman = to_roman(year) if year else ""
    semesterRoman = to_roman(semester) if semester else ""
    
    # Branch name display mapping
    branch_map = {
        'AI_DS': 'ARTIFICIAL INTELLIGENCE AND DATA SCIENCE',
        'AIDS': 'ARTIFICIAL INTELLIGENCE AND DATA SCIENCE',
        'CSE': 'COMPUTER SCIENCE AND ENGINEERING',
        'IT': 'INFORMATION TECHNOLOGY',
        'ECE': 'ELECTRONICS AND COMMUNICATION ENGINEERING',
        'EEE': 'ELECTRICAL AND ELECTRONICS ENGINEERING',
        'MECH': 'MECHANICAL ENGINEERING',
        'CIVIL': 'CIVIL ENGINEERING'
    }
    dept_upper = departmentName.upper()
    for key, val in branch_map.items():
        if key in dept_upper or val in dept_upper:
            departmentName = val
            break
    
    # Academic Year calculation
    from datetime import datetime
    now = datetime.now()
    if now.month >= 6:
        academicYear_start = now.year
    else:
        academicYear_start = now.year - 1
    academicYear = f"{academicYear_start} – {str(academicYear_start + 1)[2:]}"
    
    # Batch calculation (Batch starts when they entered 1st year)
    try:
        y_val = int(year) if year else 1
    except:
        y_val = 1
    batch_start = academicYear_start - y_val + 1
    batch_end = batch_start + 4
    batch_str = f"{batch_start} – {batch_end}"
    
    # 2. Parse generated content markdown
    units_data = parse_beyond_markdown(content)
    
    # 3. Open template
    doc = docx.Document(template_path)
    
    # Replace global header placeholders
    replace_text_in_runs(doc.paragraphs, "{{DEPARTMENT}}", departmentName.upper())
    replace_text_in_runs(doc.paragraphs, "{{ACADEMIC_YEAR}}", academicYear)
    
    # Replace table placeholders (Table 0 is metadata table)
    if len(doc.tables) > 0:
        meta_table = doc.tables[0]
        for row in meta_table.rows:
            for cell in row.cells:
                if "{{STAFF_NAME}}" in cell.text:
                    cell.text = cell.text.replace("{{STAFF_NAME}}", staffName)
                elif "{{SUBJECT_CODE}}" in cell.text:
                    cell.text = cell.text.replace("{{SUBJECT_CODE}}", subjectCode).replace("{{SUBJECT_NAME}}", subjectName)
                elif "{{YEAR}}" in cell.text:
                    cell.text = cell.text.replace("{{YEAR}}", yearRoman).replace("{{SEMESTER}}", semesterRoman)
                elif "{{BATCH}}" in cell.text:
                    cell.text = cell.text.replace("{{BATCH}}", batch_str)
                    
                # Fix font style for all cells in metadata table
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(11)
                        
    # 4. Find dynamic content placeholder paragraph
    placeholder_p = None
    for p in doc.paragraphs:
        if "{{CONTENT_BEYOND_SYLLABUS_PLACEHOLDER}}" in p.text:
            placeholder_p = p
            break
            
    if placeholder_p:
        # We will insert paragraphs before/after this placeholder, and then remove the placeholder paragraph
        # Create standard roman numeral mapper for units
        roman_nums = ["I", "II", "III", "IV", "V"]
        
        for idx, u_data in enumerate(units_data):
            roman_lbl = roman_nums[idx] if idx < len(roman_nums) else str(idx + 1)
            unit_title = u_data["title"].upper()
            
            # Add Unit Header paragraph
            p_head = doc.add_paragraph()
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_head.paragraph_format.space_before = Pt(12)
            p_head.paragraph_format.space_after = Pt(4)
            p_head.paragraph_format.keep_with_next = True
            
            run_head = p_head.add_run(f"UNIT {roman_lbl} – {unit_title}")
            run_head.font.name = 'Times New Roman'
            run_head.font.size = Pt(11)
            run_head.bold = True
            
            # Add Subheader
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_sub.paragraph_format.space_before = Pt(4)
            p_sub.paragraph_format.space_after = Pt(4)
            p_sub.paragraph_format.keep_with_next = True
            
            run_sub = p_sub.add_run("Beyond-the-Syllabus Topics:")
            run_sub.font.name = 'Times New Roman'
            run_sub.font.size = Pt(11)
            run_sub.bold = True
            
            # Add Topics
            for t_idx, topic in enumerate(u_data["topics"]):
                p_topic = doc.add_paragraph()
                p_topic.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_topic.paragraph_format.space_before = Pt(2)
                p_topic.paragraph_format.space_after = Pt(2)
                p_topic.paragraph_format.line_spacing = 1.15
                
                # Part 1: Bold number and topic name
                # E.g. "1. Automated EDA Tools "
                num_lbl = topic["num"] if topic["num"] else str(idx * 5 + t_idx + 1)
                run_topic_num = p_topic.add_run(f"{num_lbl}. {topic['name']} ")
                run_topic_num.font.name = 'Times New Roman'
                run_topic_num.font.size = Pt(11)
                run_topic_num.bold = True
                
                # Part 2: Dash separator
                run_dash = p_topic.add_run("– ")
                run_dash.font.name = 'Times New Roman'
                run_dash.font.size = Pt(11)
                
                # Part 3: Normal Description
                run_desc = p_topic.add_run(topic["desc"])
                run_desc.font.name = 'Times New Roman'
                run_desc.font.size = Pt(11)
                
            # Spacer between units
            p_unit_space = doc.add_paragraph()
            p_unit_space.paragraph_format.space_before = Pt(0)
            p_unit_space.paragraph_format.space_after = Pt(6)
            
        # Move all added paragraphs to replace the placeholder
        # In python-docx, new paragraphs are added at the end of the document, so we can re-order XML elements
        # to place them right before the placeholder_p!
        body = doc.element.body
        placeholder_idx = body.index(placeholder_p._element)
        
        # The added elements are at the end, so we can pop and insert them
        added_count = 0
        for u_data in units_data:
            # unit header + sub header + 5 topics + spacer
            added_count += 3 + len(u_data["topics"])
            
        # Move them
        elements_to_move = list(body)[-added_count:]
        for el in elements_to_move:
            body.insert(placeholder_idx, el)
            placeholder_idx += 1
            
        # Remove placeholder
        body.remove(placeholder_p._element)
        
    # Save output docx
    doc.save(output_path)
    
    # Clean up temporary input file
    try:
        os.remove(input_path)
    except:
        pass
        
    print("SUCCESS")

if __name__ == '__main__':
    main()
