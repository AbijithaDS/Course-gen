import sys
import os
import json
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def to_roman(num):
    try:
        n = int(num)
        if n <= 0:
            return str(num)
        roman_map = [
            (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
        ]
        result = ''
        for value, symbol in roman_map:
            while n >= value:
                result += symbol
                n -= value
        return result
    except:
        return str(num)

def parse_hots_markdown(md_text):
    units = []
    # Match headers like "### Unit I: Title" or "### Unit 1: Title"
    matches = list(re.finditer(r"(?:^|\n)\s*###+\s*(Unit\s+[IVX\d]+[:\-\s].*?)(?=\n|$)", md_text, re.IGNORECASE))
    
    for i, match in enumerate(matches):
        header = match.group(1).strip()
        start = match.end()
        end = matches[i+1].start() if i + 1 < len(matches) else len(md_text)
        unit_body = md_text[start:end].strip()
        
        # Extract lines
        lines = unit_body.split('\n')
        questions = []
        for line in lines:
            line_str = line.strip()
            # Match line starts with numbers, e.g. "1. Question text"
            m = re.match(r"^\s*(\d+)[\.\)]\s*(.*)", line_str)
            if m:
                q_num = m.group(1).strip()
                q_text = m.group(2).strip()
                
                # Strip markdown asterisks or underscores
                q_text = re.sub(r"\*+|_+", "", q_text).strip()
                
                questions.append({
                    "num": q_num,
                    "text": q_text
                })
                
        # Clean header title
        title_parts = header.split(':', 1)
        unit_title = title_parts[1].strip() if len(title_parts) > 1 else header
        unit_title = re.sub(r"\*+|_+", "", unit_title).strip()
        
        units.append({
            "title": unit_title,
            "questions": questions
        })
        
    return units

def replace_text_in_runs(paragraphs, placeholder, replacement):
    for p in paragraphs:
        if placeholder in p.text:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)

def main():
    if len(sys.argv) < 4:
        print("Usage: python hotsGenerator.py <templatePath> <inputPath> <outputPath>")
        sys.exit(1)
        
    template_path = sys.argv[1]
    input_path = sys.argv[2]
    output_path = sys.argv[3]
    
    # 1. Load dynamic input payload
    with open(input_path, 'r', encoding='utf-8') as f:
        payload = json.load(f)
        
    subjectCode = payload.get("subjectCode", "")
    subjectName = payload.get("subjectName", "Academic Subject")
    staffName = payload.get("staffName", "Faculty member")
    departmentName = payload.get("departmentName", "Academic Department")
    year = payload.get("year", "")
    semester = payload.get("semester", "")
    regulation = payload.get("regulation", "")
    content = payload.get("content", "")
    
    # Format metadata
    yearRoman = to_roman(year) if year else ""
    semesterRoman = to_roman(semester) if semester else ""
    
    # Branch name display mapping
    branch_map = {
        'AI_DS': 'ARTIFICIAL INTELLIGENCE AND DATA SCIENCE',
        'AIDS': 'ARTIFICIAL INTELLIGENCE AND DATA SCIENCE',
        'CSE': 'COMPUTER SCIENCE AND ENGINEERING',
        'IT': 'INFORMATION TECHNOLOGY',
        'ECE': 'ELECTRONICS AND COMMUNICATION ENGINEERING',
        'EEE': 'ELECTRICAL AND ELECTRONICS ENGINEERING',
        'MECH': 'MECHANICAL ENGINEERING',
        'CIVIL': 'CIVIL ENGINEERING'
    }
    dept_upper = departmentName.upper()
    for key, val in branch_map.items():
        if key in dept_upper or val in dept_upper:
            departmentName = val
            break
    
    # Academic Year calculation
    from datetime import datetime
    now = datetime.now()
    if now.month >= 6:
        academicYear_start = now.year
    else:
        academicYear_start = now.year - 1
    academicYear = f"{academicYear_start} – {str(academicYear_start + 1)[2:]}"
    
    # Batch calculation (Batch starts when they entered 1st year)
    try:
        y_val = int(year) if year else 1
    except:
        y_val = 1
    batch_start = academicYear_start - y_val + 1
    batch_end = batch_start + 4
    batch_str = f"{batch_start} – {batch_end}"
    
    # 2. Parse generated content markdown
    units_data = parse_hots_markdown(content)
    
    # 3. Open template
    doc = docx.Document(template_path)
    
    # Replace global header placeholders
    replace_text_in_runs(doc.paragraphs, "{{DEPARTMENT}}", departmentName.upper())
    replace_text_in_runs(doc.paragraphs, "{{ACADEMIC_YEAR}}", academicYear)
    
    # Replace table placeholders (check all tables in document)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{STAFF_NAME}}" in cell.text:
                    cell.text = cell.text.replace("{{STAFF_NAME}}", staffName)
                if "{{SUBJECT_CODE}}" in cell.text:
                    cell.text = cell.text.replace("{{SUBJECT_CODE}}", subjectCode).replace("{{SUBJECT_NAME}}", subjectName)
                if "{{YEAR}}" in cell.text:
                    cell.text = cell.text.replace("{{YEAR}}", yearRoman).replace("{{SEMESTER}}", semesterRoman)
                if "{{BATCH}}" in cell.text:
                    cell.text = cell.text.replace("{{BATCH}}", batch_str)
                if "{{REGULATION}}" in cell.text:
                    cell.text = cell.text.replace("{{REGULATION}}", regulation)
                    
                # Fix font style for all cells in metadata table
                if any(x in cell.text for x in [staffName, subjectCode, yearRoman, batch_str, regulation]):
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(11)
                        
    # 4. Find dynamic content placeholder paragraph
    placeholder_p = None
    for p in doc.paragraphs:
        if "{{HOTS_QUESTIONS_PLACEHOLDER}}" in p.text:
            placeholder_p = p
            break
            
    if placeholder_p:
        roman_nums = ["I", "II", "III", "IV", "V"]
        
        for idx, u_data in enumerate(units_data):
            roman_lbl = roman_nums[idx] if idx < len(roman_nums) else str(idx + 1)
            unit_title = u_data["title"].upper()
            
            # Add Unit Header paragraph
            p_head = placeholder_p.insert_paragraph_before()
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_head.paragraph_format.space_before = Pt(12)
            p_head.paragraph_format.space_after = Pt(4)
            p_head.paragraph_format.keep_with_next = True
            
            run_head = p_head.add_run(f"UNIT {roman_lbl} – {unit_title}")
            run_head.font.name = 'Times New Roman'
            run_head.font.size = Pt(11)
            run_head.bold = True
            
            # Add Questions
            for q_idx, q in enumerate(u_data["questions"]):
                p_q = placeholder_p.insert_paragraph_before()
                p_q.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_q.paragraph_format.space_before = Pt(2)
                p_q.paragraph_format.space_after = Pt(12)  # Give one row space after question end
                p_q.paragraph_format.line_spacing = 1.15
                
                # Number and question text
                num_lbl = q["num"] if q["num"] else str(q_idx + 1)
                run_text = p_q.add_run(f"{num_lbl}. {q['text']}")
                run_text.font.name = 'Times New Roman'
                run_text.font.size = Pt(11)
                
            # Spacer between units
            p_unit_space = placeholder_p.insert_paragraph_before()
            p_unit_space.paragraph_format.space_before = Pt(0)
            p_unit_space.paragraph_format.space_after = Pt(6)
            
        # Delete the placeholder paragraph
        p_element = placeholder_p._element
        p_element.getparent().remove(p_element)
        
    # Save output docx
    doc.save(output_path)
    
    # Clean up temporary input file
    try:
        os.remove(input_path)
    except:
        pass
        
    print("SUCCESS")

if __name__ == '__main__':
    main()
