import sys
import os
import json
import re
import docx

def to_roman(num):
    try:
        n = int(num)
        if n <= 0:
            return str(num)
        roman_map = [
            (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
        ]
        result = ''
        for value, symbol in roman_map:
            while n >= value:
                result += symbol
                n -= value
        return result
    except:
        return str(num)

def parse_qbank_markdown(md_text):
    units = []
    # Match headers like "### Unit I: Title"
    matches = list(re.finditer(r"(?:^|\n)\s*###+\s*(Unit\s+[IVX\d]+[:\-\s].*?)(?=\n|$)", md_text, re.IGNORECASE))
    
    for i, match in enumerate(matches):
        header = match.group(1).strip()
        start = match.end()
        end = matches[i+1].start() if i + 1 < len(matches) else len(md_text)
        unit_body = md_text[start:end]
        
        two_mark_section = ""
        sixteen_mark_section = ""
        
        two_mark_match = re.search(r"(?:2\-Mark|2\s+Mark|Short\s+Answer|Part\s+A)", unit_body, re.IGNORECASE)
        sixteen_mark_match = re.search(r"(?:16\-Mark|16\s+Mark|Long\s+Answer|Part\s+B)", unit_body, re.IGNORECASE)
        
        if two_mark_match and sixteen_mark_match:
            if two_mark_match.start() < sixteen_mark_match.start():
                two_mark_section = unit_body[two_mark_match.end():sixteen_mark_match.start()]
                sixteen_mark_section = unit_body[sixteen_mark_match.end():]
            else:
                sixteen_mark_section = unit_body[sixteen_mark_match.end():two_mark_match.start()]
                two_mark_section = unit_body[two_mark_match.end():]
        elif two_mark_match:
            two_mark_section = unit_body[two_mark_match.end():]
        elif sixteen_mark_match:
            sixteen_mark_section = unit_body[sixteen_mark_match.end():]
        else:
            two_mark_section = unit_body
            
        def extract_questions(section_text):
            lines = section_text.split('\n')
            qs = []
            for line in lines:
                line_str = line.strip()
                # Match line starts with numbers, e.g. "1. Question"
                m = re.match(r"^\s*\d+[\.\)]\s*(.+)", line_str)
                if m:
                    q_text = m.group(1).strip()
                    co = ""
                    kx = ""
                    
                    co_k_match = re.search(r"\(\s*(CO[1-5]|K[1-6])\s*(?:,\s*(CO[1-5]|K[1-6]))?\s*\)", q_text, re.IGNORECASE)
                    if co_k_match:
                        for g in co_k_match.groups():
                            if g:
                                if g.upper().startswith("CO"):
                                    co = g.upper()
                                elif g.upper().startswith("K"):
                                    kx = g.upper()
                        q_text = q_text.replace(co_k_match.group(0), "")
                        
                    single_matches = re.findall(r"\(\s*(CO[1-5]|K[1-6])\s*\)", q_text, re.IGNORECASE)
                    for sm in single_matches:
                        if sm.upper().startswith("CO"):
                            co = sm.upper()
                        elif sm.upper().startswith("K"):
                            kx = sm.upper()
                        q_text = re.sub(r"\(\s*" + sm + r"\s*\)", "", q_text, flags=re.IGNORECASE)
                        
                    q_text = re.sub(r"\*+|_+", "", q_text).strip()
                    qs.append({"text": q_text, "co": co, "kx": kx})
            return qs
            
        two_mark_qs = extract_questions(two_mark_section)
        sixteen_mark_qs = extract_questions(sixteen_mark_section)
        
        units.append({
            "header": header,
            "two_mark": two_mark_qs,
            "sixteen_mark": sixteen_mark_qs
        })
        
    return units

def set_cell_text_preserving_formatting(cell, text):
    if len(cell.paragraphs) > 0:
        p = cell.paragraphs[0]
        if len(p.runs) > 0:
            p.runs[0].text = text
            for r in list(p.runs[1:]):
                p._element.remove(r._r)
        else:
            p.add_run(text)
    else:
        cell.text = text

def set_cell_bold_text(cell, text):
    if len(cell.paragraphs) > 0:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.text = ""
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'

def set_cell_bold_centered_text(cell, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if len(cell.paragraphs) > 0:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'

def main():
    if len(sys.argv) < 4:
        print("Usage: python qbankGenerator.py <templatePath> <inputPath> <outputPath>")
        sys.exit(1)
        
    template_path = sys.argv[1]
    input_path = sys.argv[2]
    output_path = sys.argv[3]
    
    # 1. Load dynamic input payload
    with open(input_path, 'r', encoding='utf-8') as f:
        payload = json.load(f)
        
    subjectCode = payload.get("subjectCode", "")
    subjectName = payload.get("subjectName", "Academic Subject")
    staffName = payload.get("staffName", "Faculty member")
    departmentName = payload.get("departmentName", "Academic Department")
    year = payload.get("year", "")
    semester = payload.get("semester", "")
    regulation = payload.get("regulation", "")
    content = payload.get("content", "")
    
    # Format metadata
    yearRoman = to_roman(year) if year else ""
    semesterRoman = to_roman(semester) if semester else ""
    
    # Academic Year calculation
    from datetime import datetime
    now = datetime.now()
    if now.month >= 6:
        academicYear = f"{now.year}-{now.year + 1}"
    else:
        academicYear = f"{now.year - 1}-{now.year}"
        
    # 2. Parse generated content markdown
    units_data = parse_qbank_markdown(content)
    
    # 3. Open template
    doc = docx.Document(template_path)
    
    # 4. Fill Header tables (Table 0, 2, 5, 7)
    header_tables = [0, 2, 5, 7]
    for idx in header_tables:
        if idx >= len(doc.tables):
            continue
        t = doc.tables[idx]
        
        # Row 0: Subject details (Value goes to Cell 1, Label in Cell 0 is preserved)
        if len(t.rows) > 0 and len(t.rows[0].cells) > 1:
            set_cell_text_preserving_formatting(t.rows[0].cells[1], f"{subjectCode} - {subjectName}")
            
        # Row 1: Faculty details (Value goes to Cell 1, Label in Cell 0 is preserved)
        if len(t.rows) > 1 and len(t.rows[1].cells) > 1:
            set_cell_text_preserving_formatting(t.rows[1].cells[1], f"{staffName}, Assistant Professor, {departmentName}")
            
        # Row 2: Year, Semester, Academic Year
        if len(t.rows) > 2:
            for cell in t.rows[2].cells:
                txt = cell.text.strip()
                if 'YEAR' in txt:
                    set_cell_bold_text(cell, f"YEAR : {yearRoman}")
                elif 'Semester' in txt:
                    set_cell_bold_text(cell, f"Semester : {semesterRoman}")
                elif 'Academic Year' in txt:
                    set_cell_bold_text(cell, f"Academic Year: {academicYear}")
                    
    # 5. Fill Q-tables (Table 1, 3, 4, 6, 8)
    q_tables = [1, 3, 4, 6, 8]
    for u_idx, t_idx in enumerate(q_tables):
        if t_idx >= len(doc.tables):
            continue
        if u_idx >= len(units_data):
            break
            
        unit_data = units_data[u_idx]
        t = doc.tables[t_idx]
        
        # 1. Update Unit Title in Row 0
        header = unit_data["header"]
        title_parts = header.split(':', 1)
        unit_title = title_parts[1].strip() if len(title_parts) > 1 else ""
        
        unit_label = f"UNIT {['I','II','III','IV','V'][u_idx]}"
        row0_text = f"{unit_label}: {unit_title.upper()}"
        for cell in t.rows[0].cells:
            set_cell_bold_centered_text(cell, row0_text)
            
        # 2. Discover columns
        header_row = t.rows[1]
        text_indices = []
        co_indices = []
        k_indices = []
        
        for c_idx, cell in enumerate(header_row.cells):
            txt = cell.text.strip().upper()
            if 'MARK' in txt or '5 X 2' in txt or '4 X 16' in txt:
                text_indices.append(c_idx)
            elif 'CO' in txt:
                co_indices.append(c_idx)
            elif 'K' in txt:
                k_indices.append(c_idx)
                
        # 3. Populate 2-mark questions (Rows 2 to 6)
        two_mark_qs = unit_data["two_mark"]
        for q_idx in range(5):
            row_idx = 2 + q_idx
            if row_idx >= len(t.rows):
                continue
            row = t.rows[row_idx]
            q_data = two_mark_qs[q_idx] if q_idx < len(two_mark_qs) else {"text": "", "co": "", "kx": ""}
            
            for col_idx in text_indices:
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = q_data["text"]
                    
            co_val = q_data["co"] if q_data["co"] else f"CO{u_idx+1}"
            for col_idx in co_indices:
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = co_val
                    
            k_val = q_data["kx"] if q_data["kx"] else "K2"
            for col_idx in k_indices:
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = k_val
                    
        # 4. Populate 16-mark questions (Rows 8 to 11)
        sixteen_mark_qs = unit_data["sixteen_mark"]
        for q_idx in range(4):
            row_idx = 8 + q_idx
            if row_idx >= len(t.rows):
                continue
            row = t.rows[row_idx]
            q_data = sixteen_mark_qs[q_idx] if q_idx < len(sixteen_mark_qs) else {"text": "", "co": "", "kx": ""}
            
            for col_idx in text_indices:
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = q_data["text"]
                    
            co_val = q_data["co"] if q_data["co"] else f"CO{u_idx+1}"
            for col_idx in co_indices:
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = co_val
                    
            k_val = q_data["kx"] if q_data["kx"] else "K3"
            for col_idx in k_indices:
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = k_val
                    
    # Save output docx
    doc.save(output_path)
    
    # 6. Clean up temporary input file
    try:
        os.remove(input_path)
    except:
        pass
        
    print("SUCCESS")

if __name__ == '__main__':
    main()
